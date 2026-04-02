from docx import Document
import zipfile
import base64
import os
import json
from lxml import etree
import warnings
import officemath2latex

try:
    import latex2mathml.converter
    LATEX2MATHML_AVAILABLE = True
    print("✓ latex2mathml загружен успешно")
except ImportError:
    LATEX2MATHML_AVAILABLE = False
    print("⚠️ latex2mathml не установлен")
warnings.filterwarnings('ignore', category=UserWarning)

# === ИМПОРТ КОНВЕРТАЦИИ WMF ===
from .db import convert_wmf_to_png
print("=== docx_parser.py ЗАГРУЖЕН УСПЕШНО ===")
print("DEBUG: convert_wmf_to_png успешно импортирован из db")

def latex_to_mathml(latex_string):
    if not LATEX2MATHML_AVAILABLE:
        return None
    try:
        mathml = latex2mathml.converter.convert(latex_string)
        print(f" ✓ MathML создан (длина: {len(mathml)})")
        return mathml
    except Exception as e:
        print(f" ⚠️ Ошибка LaTeX->MathML: {e}")
        return None

def extract_omml_native(omath_element):
    try:
        xml_str = etree.tostring(omath_element, encoding='unicode', pretty_print=True)
        return xml_str
    except Exception as e:
        print(f"⚠️ Ошибка извлечения OMML: {e}")
        return None

def omml_to_latex(omml_xml):
    try:
        latex = officemath2latex.process_math_string(omml_xml)
        if latex:
            print(f" ✓ LaTeX создан (длина: {len(latex)})")
            return latex
        return None
    except Exception as e:
        print(f"⚠️ Ошибка OMML->LaTeX: {e}")
        return None

def build_rid_to_media_map(file_path, images_data):
    rid_to_media = {}
    with zipfile.ZipFile(file_path, 'r') as z:
        try:
            rels_content = z.read('word/_rels/document.xml.rels')
            rels_root = etree.fromstring(rels_content)
            for rel in rels_root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship'):
                rel_id = rel.get('Id')
                rel_target = rel.get('Target')
                if rel_target and 'media/' in rel_target:
                    full_path = f"word/{rel_target}"
                    if full_path in images_data:
                        rid_to_media[rel_id] = full_path
        except Exception as e:
            print(f"⚠️ Ошибка relationships: {e}")
    return rid_to_media

def extract_images_and_formulas(file_path):
    images_data = {}
    has_wmf = False
    with zipfile.ZipFile(file_path, 'r') as docx_zip:
        for file_name in docx_zip.namelist():
            if file_name.startswith('word/media/'):
                file_data = docx_zip.read(file_name)
                base64_data = base64.b64encode(file_data).decode('utf-8')
                file_ext = os.path.splitext(file_name)[1].lower()
                if file_ext == '.wmf':
                    has_wmf = True
                mime_type = {'.png':'image/png','.jpg':'image/jpeg','.jpeg':'image/jpeg','.gif':'image/gif','.bmp':'image/bmp','.wmf':'image/x-wmf','.emf':'image/x-emf'}.get(file_ext, 'application/octet-stream')
                images_data[file_name] = {'base64': base64_data, 'mime_type': mime_type, 'file_name': os.path.basename(file_name), 'extension': file_ext}
    if has_wmf:
        print("⚠️ ОБНАРУЖЕН WMF — будет конвертация")
    return images_data, has_wmf

def process_element_media(element, media_counter, formula_counter, images_data, media_registry, formulas_registry, rid_to_media):
    cell_media = []
    # Inline shapes
    for shape in element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'):
        for blip in shape.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
            if media_counter <= len(images_data):
                img_info = list(images_data.values())[media_counter - 1]
                media_label = f"[media{media_counter}]"
                if img_info['extension'] == '.wmf':
                    print(f"DEBUG: === НАЙДЕН WMF {media_label} — конвертация ===")
                    png_base64, new_mime = convert_wmf_to_png(img_info['base64'])
                    if new_mime == 'image/png':
                        media_registry[media_label] = {'base64': png_base64, 'mime_type': 'image/png', 'file_name': img_info['file_name'].replace('.wmf', '.png'), 'original_format': 'wmf'}
                        print(f"DEBUG: WMF {media_label} → PNG УСПЕШНО")
                    else:
                        media_registry[media_label] = {'base64': img_info['base64'], 'mime_type': img_info['mime_type'], 'file_name': img_info['file_name'], 'conversion_failed': True}
                else:
                    media_registry[media_label] = {'base64': img_info['base64'], 'mime_type': img_info['mime_type'], 'file_name': img_info['file_name']}
                cell_media.append(media_label)
                media_counter += 1
    # VML shapes
    for vml_shape in element.findall('.//{urn:schemas-microsoft-com:vml}shape'):
        for imgdata in vml_shape.findall('.//{urn:schemas-microsoft-com:vml}imagedata'):
            rid = imgdata.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
            if rid and rid in rid_to_media:
                media_path = rid_to_media[rid]
                if media_path in images_data:
                    img_info = images_data[media_path]
                    media_label = f"[media{media_counter}]"
                    if img_info['extension'] == '.wmf':
                        print(f"DEBUG: === VML WMF {media_label} — конвертация ===")
                        png_base64, new_mime = convert_wmf_to_png(img_info['base64'])
                        if new_mime == 'image/png':
                            media_registry[media_label] = {'base64': png_base64, 'mime_type': 'image/png', 'file_name': img_info['file_name'].replace('.wmf', '.png'), 'original_format': 'wmf', 'source': 'vml'}
                        else:
                            media_registry[media_label] = {'base64': img_info['base64'], 'mime_type': img_info['mime_type'], 'file_name': img_info['file_name'], 'conversion_failed': True, 'source': 'vml'}
                    else:
                        media_registry[media_label] = {'base64': img_info['base64'], 'mime_type': img_info['mime_type'], 'file_name': img_info['file_name'], 'source': 'vml'}
                    cell_media.append(media_label)
                    media_counter += 1
    # Anchored images
    anchors = element.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}anchor')
    for anchor in anchors:
        for shape in anchor.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/picture}pic'):
            for blip in shape.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip'):
                if media_counter <= len(images_data):
                    img_info = list(images_data.values())[media_counter - 1]
                    media_label = f"[media{media_counter}]"
                    media_registry[media_label] = {'base64': img_info['base64'], 'mime_type': img_info['mime_type'], 'file_name': img_info['file_name']}
                    cell_media.append(media_label)
                    media_counter += 1
    # Формулы OMML
    omath_elements = element.findall('.//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath')
    for omath in omath_elements:
        omml_xml = extract_omml_native(omath)
        if omml_xml:
            formula_label = f"[formula{formula_counter}]"
            latex = omml_to_latex(omml_xml)
            if latex:
                mathml = latex_to_mathml(latex)
                formulas_registry[formula_label] = {'latex': latex, 'mathml': mathml or '', 'formula_preview': ''.join(omath.itertext())[:50], 'type': 'latex'}
            else:
                omml_base64 = base64.b64encode(omml_xml.encode('utf-8')).decode('utf-8')
                formulas_registry[formula_label] = {'omml_xml': omml_base64, 'type': 'omml'}
            cell_media.append(formula_label)
            formula_counter += 1
    return cell_media, media_counter, formula_counter

def parse_docx_to_json(file_path, output_json_path='output.json'):
    print("DEBUG: === НАЧАЛСЯ ПАРСИНГ DOCX ===")
    images_data, has_wmf = extract_images_and_formulas(file_path)
    print(f"DEBUG: Найдено изображений: {len(images_data)}, WMF: {has_wmf}")
    rid_to_media = build_rid_to_media_map(file_path, images_data)
    media_counter = 1
    media_registry = {}
    formula_counter = 1
    formulas_registry = {}
    doc = Document(file_path)
    if not doc.tables:
        print("❌ Таблицы не найдены")
        return {"questions": []}
    table = doc.tables[0]
    all_rows = []
    for row in table.rows:
        row_data = []
        for cell in row.cells:
            cell_text = cell.text.strip()
            cell_media, media_counter, formula_counter = process_element_media(cell._element, media_counter, formula_counter, images_data, media_registry, formulas_registry, rid_to_media)
            # txbxContent
            for txbx in cell._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}txbxContent'):
                txbx_text_parts = [t.text for t in txbx.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t') if t.text]
                if txbx_text_parts:
                    cell_text += " " + "".join(txbx_text_parts)
                txbx_media, media_counter, formula_counter = process_element_media(txbx, media_counter, formula_counter, images_data, media_registry, formulas_registry, rid_to_media)
                cell_media.extend(txbx_media)
            row_data.append({'text': cell_text.strip(), 'media': cell_media})
        all_rows.append(row_data)
    print(f"DEBUG: Обработано строк: {len(all_rows)}")

    # Формирование вопросов
    questions_list = []
    if len(all_rows) > 1:
        header_row = all_rows[0]
        question_column = None
        correct_answer_columns = []
        wrong_answer_columns = []
        labels = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']
        for index, cell_data in enumerate(header_row):
            cell_value = cell_data['text']
            if "Вопрос" in cell_value or cell_value == "Вопросы":
                question_column = index
            elif "Правильный" in cell_value or "(Правильный)" in cell_value:
                correct_answer_columns.append(index)
            elif any(letter in cell_value for letter in labels):
                if index not in correct_answer_columns:
                    wrong_answer_columns.append(index)
        for row_index in range(1, len(all_rows)):
            current_row = all_rows[row_index]
            question_text = ""
            if question_column is not None and question_column < len(current_row):
                question_data = current_row[question_column]
                question_text = question_data['text']
                for media_label in question_data['media']:
                    question_text += f" {media_label}"
            all_answers = []
            correct_indices = []
            for col_index in correct_answer_columns:
                if col_index < len(current_row):
                    answer_data = current_row[col_index]
                    answer_text = answer_data['text']
                    for media_label in answer_data['media']:
                        answer_text += f" {media_label}"
                    correct_indices.append(len(all_answers))
                    all_answers.append(answer_text)
            for col_index in wrong_answer_columns:
                if col_index < len(current_row):
                    answer_data = current_row[col_index]
                    answer_text = answer_data['text']
                    for media_label in answer_data['media']:
                        answer_text += f" {media_label}"
                    all_answers.append(answer_text)
            question_media = {}
            question_formulas = {}
            full_text = question_text + " " + " ".join(all_answers)
            for label in media_registry:
                if label in full_text:
                    question_media[label] = media_registry[label]
            for label in formulas_registry:
                if label in full_text:
                    question_formulas[label] = formulas_registry[label]
            question_obj = {
                "id": f"q{row_index}",
                "question": question_text,
                "answers": all_answers,
                "correct_indices": correct_indices,
                "media": question_media,
                "formulas": question_formulas
            }
            questions_list.append(question_obj)
    result_json = {"questions": questions_list}
    with open(output_json_path, 'w', encoding='utf-8') as f:
        json.dump(result_json, f, ensure_ascii=False, indent=2)
    print(f"DEBUG: JSON сохранён → вопросов найдено: {len(questions_list)}")
    return result_json
